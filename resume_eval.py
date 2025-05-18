import streamlit as st
from io import BytesIO
from docx import Document
from datetime import date

st.set_page_config(page_title="Resume Builder", layout="centered")

st.title("📝 Resume Builder")
st.markdown("Fill in the details below to generate your resume.")

# User Input Fields
with st.form("resume_form"):
    st.header("Personal Information")
    name = st.text_input("Full Name")
    email = st.text_input("Email")
    phone = st.text_input("Phone Number")
    linkedin = st.text_input("LinkedIn URL")
    github = st.text_input("GitHub URL")
    summary = st.text_area("Career Summary", height=100)

    st.header("Education")
    edu_degree = st.text_input("Degree")
    edu_school = st.text_input("School/University")
    edu_year = st.text_input("Year of Graduation")

    st.header("Experience")
    exp_title = st.text_input("Job Title")
    exp_company = st.text_input("Company")
    exp_duration = st.text_input("Duration")
    exp_description = st.text_area("Description", height=100)

    st.header("Skills")
    skills = st.text_area("List your skills separated by commas")

    submitted = st.form_submit_button("Generate Resume")

# Function to generate DOCX Resume
def create_resume():
    doc = Document()
    
    doc.add_heading(name, 0)
    doc.add_paragraph(f"Email: {email} | Phone: {phone}")
    doc.add_paragraph(f"LinkedIn: {linkedin} | GitHub: {github}")
    doc.add_paragraph()
    
    doc.add_heading("Career Summary", level=1)
    doc.add_paragraph(summary)
    
    doc.add_heading("Education", level=1)
    doc.add_paragraph(f"{edu_degree}, {edu_school} ({edu_year})")
    
    doc.add_heading("Experience", level=1)
    doc.add_paragraph(f"{exp_title} - {exp_company} ({exp_duration})")
    doc.add_paragraph(exp_description)
    
    doc.add_heading("Skills", level=1)
    doc.add_paragraph(skills)

    # Save to memory buffer
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

if submitted:
    if name and email:
        resume_file = create_resume()
        st.success("✅ Resume generated successfully!")
        st.download_button(
            label="📄 Download Resume (DOCX)",
            data=resume_file,
            file_name=f"{name}_Resume.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    else:
        st.error("❌ Name and Email are required to generate the resume.")
